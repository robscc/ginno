"""Unit tests: preview payloads (paginated grids + document markdown)."""

from __future__ import annotations

from pathlib import Path

import pytest

from ginno_runtime.files import extractors as ex
from ginno_runtime.files.preview import build_csv_export, build_preview

pytestmark = pytest.mark.unit


def make_big_xlsx(p: Path, rows: int = 150) -> Path:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "data"
    ws.append(["id", "value"])
    for i in range(rows):
        ws.append([i, i * 1.5])
    ws2 = wb.create_sheet("other")
    ws2.append(["x"])
    wb.save(p)
    return p


def test_spreadsheet_preview_pagination(tmp_path):
    f = make_big_xlsx(tmp_path / "big.xlsx", rows=150)
    pv = build_preview(f, offset=100, limit=100)
    assert pv["kind"] == "spreadsheet"
    assert pv["sheet"] == "data"  # first sheet by default
    assert {s["name"] for s in pv["sheets"]} == {"data", "other"}
    assert pv["total_rows"] == 150
    assert len(pv["rows"]) == 50  # only 50 left after offset 100
    # cells are stringified; xlsx stores numbers as float, so accept "100"/"100.0"
    assert float(pv["rows"][0][0]) == 100
    assert pv["rows"][0][1] == "150.0"
    assert {c["name"] for c in pv["columns"]} == {"id", "value"}


def test_spreadsheet_preview_sheet_select_and_clamp(tmp_path):
    f = make_big_xlsx(tmp_path / "big.xlsx")
    pv = build_preview(f, sheet="other")
    assert pv["sheet"] == "other"
    assert pv["total_rows"] == 0
    # bogus sheet falls back to the first
    pv2 = build_preview(f, sheet="nope")
    assert pv2["sheet"] == "data"
    # limit clamps into [1, 500]
    pv3 = build_preview(f, limit=99999)
    assert pv3["limit"] == 500
    pv4 = build_preview(f, limit=0)
    assert pv4["limit"] == 1


def test_csv_preview(tmp_path):
    f = tmp_path / "t.csv"
    f.write_text("a,b\n1,2\n", encoding="utf-8")
    pv = build_preview(f)
    assert pv["kind"] == "table"
    assert pv["rows"] == [["1", "2"]]
    assert pv["total_rows"] == 1


def test_document_preview_markdown(tmp_path):
    from docx import Document

    f = tmp_path / "d.docx"
    doc = Document()
    doc.add_heading("标题", level=1)
    doc.add_paragraph("正文段落")
    doc.save(f)
    pv = build_preview(f)
    assert pv["kind"] == "document"
    assert "正文段落" in pv["markdown"]
    assert "metadata" in pv


def test_unsupported_preview(tmp_path):
    f = tmp_path / "x.exe"
    f.write_bytes(b"MZ")
    with pytest.raises(ex.UnsupportedFormat):
        build_preview(f)


# --------------------------------------------------------------------------
# CSV export
# --------------------------------------------------------------------------

def test_csv_export_roundtrip(tmp_path):
    f = tmp_path / "数据.csv"
    f.write_text("名称,金额\n甲,1.5\n乙,2\n", encoding="utf-8")
    name, data = build_csv_export(f)
    assert name == "数据.csv"
    text = data.decode("utf-8-sig")  # BOM so Excel opens Chinese correctly
    assert data.startswith(b"\xef\xbb\xbf")
    assert "名称,金额" in text and "甲,1.5" in text


def test_csv_export_xlsx_sheet_select(tmp_path):
    f = make_big_xlsx(tmp_path / "big.xlsx")  # sheets: data (150), other (0)
    name, data = build_csv_export(f)
    assert name == "big-data.csv"  # multi-sheet → suffixed with active sheet
    lines = data.decode("utf-8-sig").strip().splitlines()
    assert lines[0] == "id,value"
    assert len(lines) == 151  # header + 150 rows
    name2, data2 = build_csv_export(f, sheet="other")
    assert name2 == "big-other.csv"
    assert data2.decode("utf-8-sig").strip() == "x"  # header only


def test_csv_export_tsv_converts_to_commas(tmp_path):
    f = tmp_path / "t.tsv"
    f.write_text("a\tb\n1\t2\n", encoding="utf-8")
    name, data = build_csv_export(f)
    assert name == "t.csv"
    assert "a,b" in data.decode("utf-8-sig")


def test_csv_export_rejects_non_table(tmp_path):
    from docx import Document

    f = tmp_path / "d.docx"
    doc = Document()
    doc.add_paragraph("x")
    doc.save(f)
    with pytest.raises(ex.UnsupportedFormat):
        build_csv_export(f)


def test_csv_export_name_override(tmp_path):
    # uploads are stored with a uuid prefix; the display name must win
    f = tmp_path / "8a24fbf7-报表.csv"
    f.write_text("a\n1\n", encoding="utf-8")
    name, _ = build_csv_export(f, name="报表.csv")
    assert name == "报表.csv"
